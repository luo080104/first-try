# -*- coding: utf-8 -*-
"""完全重建简历——保留原表格，每块内容加上标题"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import copy

SRC_PATH = r"C:\Users\luoji\Desktop\个人资料\简历初稿_骆嘉铭(1).docx"
OUT_PATH = r"C:\Users\luoji\Desktop\个人资料\简历初稿_骆嘉铭_完整版.docx"

# ===== 内容 =====
sections = [
    {
        "title": "个人获奖经历",
        "desc": "（在校期间获得的各类奖项，如奖学金、竞赛获奖、荣誉称号等）",
        "items": [
            "2026年3月  高三三模总分650分（超一本线约120分）  校级卓越奖",
            "2025年12月  杭州市优秀学生  杭州市教育局",
            "2025年10月  英语单科状元（130/150分）  杭师大附属未来科技城学校",
            "2025-2026学年  三好学生（连续获评）  杭师大附属未来科技城学校",
            "2025-2026学年  一等奖学金  杭师大附属未来科技城学校",
            "2024-2025学年  二等奖学金  杭师大附属未来科技城学校",
            "2024年11月  奋进杯数学竞赛一等奖  余杭区教育局",
            "2025年12月  优秀寝室长  天元公学",
            "2024年10月  期中考试勇立潮头奖  杭师大附属未来科技城学校",
            "2023年11月  余杭区高中生应急救护技能比赛个人三等奖  余杭区红十字会",
        ]
    },
    {
        "title": "班干部工作经历",
        "desc": "（担任过的班委职务及主要工作内容）",
        "items": [
            "2024年9月至2026年6月  担任班级学习委员  负责收发作业、传达教务通知、组织学习帮扶小组，协助班主任管理班级学习事务。任职期间班级整体成绩稳步提升。",
            "2025年3月至2026年6月  担任寝室长  获评优秀寝室长称号，维持寝室纪律与卫生，室友成绩均进入年级前列。",
        ]
    },
    {
        "title": "志愿经历",
        "desc": "（参与过的志愿服务、公益活动等）",
        "items": [
            "2024年7月  参与社区科普志愿服务  累计12小时  面向社区居民开展人工智能与编程知识普及讲座，辅导青少年科技兴趣小组。",
            "2023年8月  余杭区红十字应急救护志愿服务  累计8小时  通过区级应急救护技能比赛检验，获个人三等奖，具备基本急救能力。",
            "2025年3月  校园开放日引导志愿服务  累计6小时  接待来访家长与学生，介绍学校办学特色与学习生活环境。",
        ]
    },
    {
        "title": "擅长技能",
        "desc": "（包括但不限于：计算机/编程能力、外语水平、文体特长、专业相关技能等）",
        "items": [
            "Python 编程（熟练）——独立开发Go购跨平台购物比价 AI Agent 项目（FastAPI + SQLite + DeepSeek大模型），支持淘宝/京东/拼多多/唯品会四平台商品搜索与智能比价。项目上线 GitHub（luo080104）并持续迭代。",
            "AI/大模型应用开发——熟练掌握 DeepSeek API、LLM 构建、Agent 设计与编排（ReAct 循环），能基于大模型实现意图解析、商品推荐、AI 购买建议等智能功能。",
            "英语 CET-4 水平——高中阶段英语长期保持 120-130/150 分数段，曾获校级英语单科状元。",
            "熟练使用 Office 办公软件（Word/Excel/PPT）。",
            "数据分析与评估基础——掌握 Excel 数据处理、SQLite 数据库增删改查，有模型评估与算法测试实习经验（金融NLP方向）。",
            "Web 前端基础（HTML/CSS/JavaScript）——独立完成Go购前端页面（PWA响应式 + ECharts数据可视化 + SSE流式渲染）。",
            "Git/GitHub 版本控制——日常使用 Git 管理个人项目，熟悉 PR/Merge/Issue 协作流程。",
        ]
    },
]

class_build_title = "对班级建设的想法"
class_build_desc = "（围绕学习氛围营造、班级凝聚力建设、班委团队协作、活动组织等方面谈谈自己的想法和建议）"
class_build = (
    "作为一名对智能制造充满热情的学生，我希望能将技术与协作理念融入班级建设中：\n\n"
    "1. 学习方面：建议建立互助学习小组，按专业课程分组，定期组织经验分享会。"
    "特别是编程和数学类课程，我可以利用自己在 Python 开发和 AI 应用方面的经验，"
    "帮助同学快速上手，实现以强带弱、共同进步。\n\n"
    "2. 专业拓展：结合智能制造专业特色，组织参观制造业企业、实验室开放日等活动，"
    "拉近课本理论与产业实践的距离。我在高二暑假曾在金融科技公司参与过模型评估实习，"
    "深知实践对学习的推动作用。\n\n"
    "3. 班级凝聚力：建议每学期组织1-2次团建活动（如运动友谊赛、编程马拉松、科创分享会等），"
    "让同学们在课堂之外建立更深厚的友谊。\n\n"
    "4. 班委协作：班委之间应定期沟通（建议双周例会），明确分工、相互补位。"
    "学习委员负责学风建设，生活委员关心同学生活，文体委员组织活动——各司其职，形成合力。\n\n"
    "5. 数字化管理：建议用在线文档维护班级通讯录与活动记录，借助技术手段提高班级管理效率。"
    "这是智能制造学生最应该擅长的方向。\n\n"
    "我相信，一个有技术底蕴、有团队凝聚力、有创新精神的班级，"
    "一定能成为智能制造专业最亮眼的存在。"
)

date_text = "2026年8月10日"

# ===== 重建文档 =====
doc = Document(SRC_PATH)

# 保存原始 header table（前5段+表格）
# 原始结构: P0=个人简历, P1=表格, P2=空行, P3-P...=内容

# 找到表格（第一个表格是个人信息表）
header_table = doc.tables[0] if doc.tables else None

# 删除从个人简历标题之后的所有内容（保留P0和表格）
# 保留前3个元素（标题+表格+空行）
body = doc.element.body
children = list(body)
# 找到所有段落和表格
elements_to_keep = []
table_count = 0
for child in children:
    tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
    if tag == 'tbl':
        table_count += 1
        if table_count <= 1:
            elements_to_keep.append(child)
    elif tag == 'p':
        # Keep first paragraph (个人简历 title) and stuff before sections
        if len(elements_to_keep) < 3:
            elements_to_keep.append(child)

# Remove all children
for child in list(body):
    body.remove(child)

# Add back kept elements
for child in elements_to_keep:
    body.append(child)

# ===== 辅助函数 =====
def add_section_title(doc, text):
    """Add a bold section title paragraph"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_desc(doc, text):
    """Add a grey description line"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_bullet(doc, text):
    """Add a bullet point"""
    p = doc.add_paragraph()
    run = p.add_run('• ' + text)
    run.font.size = Pt(10.5)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.5)
    return p

def add_body(doc, text):
    """Add a normal body paragraph"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p

# ===== 填充所有section =====
for sec in sections:
    add_section_title(doc, sec['title'])
    add_desc(doc, sec['desc'])
    for item in sec['items']:
        add_bullet(doc, item)
    doc.add_paragraph()  # spacer

# 班级建设
add_section_title(doc, class_build_title)
add_desc(doc, class_build_desc)
add_body(doc, class_build)
doc.add_paragraph()

# 日期
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run(date_text)
run.font.size = Pt(10.5)
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 保存
doc.save(OUT_PATH)
print(f'OK -> {OUT_PATH}')

# 验证
from docx import Document as DocReader
doc2 = DocReader(OUT_PATH)
print(f'Total paragraphs: {len(doc2.paragraphs)}')
for i, p in enumerate(doc2.paragraphs):
    t = p.text.strip()
    if t:
        print(f'  P{i}: {t[:100]}')
